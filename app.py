# import os
# import re
# from datetime import datetime
# from flask import Flask, render_template, request, redirect, url_for, flash,jsonify
# from flask_sqlalchemy import SQLAlchemy
# from werkzeug.utils import secure_filename
# import PyPDF2
# import docx
# import sqlite3
# import plotly.graph_objs as go
# import plotly
# import json
# from groq import Groq


# app = Flask(__name__)
# app.secret_key = 'your_secret_key_here'
# app.config['UPLOAD_FOLDER'] = 'uploads'
# app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///cv_database.db'
# app.config['SQLALCHEMY_TRACK_CHANGES'] = False
# app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB max file size

# # Ensure upload folder exists
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# db = SQLAlchemy(app)

# groq_client = Groq(
#     api_key="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"
# )



# system_prompt = """
# # Recruitment Analysis Assistant

# You are an expert AI recruitment analyst specializing in talent acquisition metrics and workforce insights. Your primary role is to analyze recruitment data and generate comprehensive, data-driven reports with strategic recommendations.

# ## Your Capabilities

# - Analyze candidate status distributions and identify bottlenecks in the recruitment pipeline
# - Generate actionable insights based on recruitment metrics
# - Provide strategic recommendations to improve hiring efficiency
# - Identify trends and patterns in recruitment data
# - Offer industry benchmarks and best practices when relevant

# ## Guidelines for Analysis

# When analyzing recruitment data:
# 1. First acknowledge the current state of the recruitment pipeline based on the provided metrics
# 2. Identify potential issues or bottlenecks in the recruitment process
# 3. Provide concrete, actionable recommendations with expected outcomes
# 4. Include quantitative targets where appropriate
# 5. Consider industry standards and best practices in your analysis
# 6. Organize insights in a clear, structured format

# ## Response Format

# Structure your reports with the following sections:
# - Executive Summary (brief overview of key findings)
# - Current Pipeline Status (detailed analysis of provided metrics)
# - Key Insights (interpretation of the data and identification of patterns)
# - Strategic Recommendations (3-5 specific, actionable items)
# - Expected Outcomes (projections for implementing recommendations)

# Remember to maintain a professional, consultative tone while providing insights that would be valuable to recruitment managers and HR executives.
# """

# #creating a connection
# def get_db_connection():
#     conn = sqlite3.connect('recruitment.db')
#     conn.row_factory = sqlite3.Row  # to get dict-like rows
#     return conn


# def get_table_schema():
   
#     conn = get_db_connection()
#     cursor = conn.cursor()
    

#     cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
#     tables = cursor.fetchall()
    
#     schema = {}
#     for table in tables:
#         table_name = table['name']
#         cursor.execute(f"PRAGMA table_info({table_name});")
#         columns = cursor.fetchall()
#         schema[table_name] = [column['name'] for column in columns]
    
#     conn.close()
#     return schema


# def natural_language_to_sql(query):
    
#     schema = get_table_schema()
    
  
#     schema_info = "Database Schema:\n"
#     for table, columns in schema.items():
#         schema_info += f"Table: {table}\n"
#         schema_info += f"Columns: {', '.join(columns)}\n\n"
    
#     # Create the prompt
#     prompt = f"""
# {schema_info}

# Based on the schema above, convert the following natural language query to a valid SQL query:
# "{query}"

# Guidelines:
# - Return ONLY the SQL query without any explanation or markdown formatting
# - Ensure the query is valid SQLite syntax
# - If aggregation is requested, include appropriate GROUP BY statements
# - If the query is asking for a visualization, include an appropriate SELECT statement for the visualization
# - Return a query that will return a reasonable number of rows (use LIMIT if necessary)
# - For time-based queries, use the appropriate SQL date functions
# - If you're not sure about a column name, make your best guess based on the schema

# SQL Query:
# """
    
#     # Generate SQL query
#     response = groq_client.chat.completions.create(
#         messages=[
#             {
#                 "role": "system",
#                 "content": "You are an expert in converting natural language to SQL queries. You carefully analyze database schemas and produce only valid SQL code."
#             },
#             {
#                 "role": "user",
#                 "content": prompt
#             }
#         ],
#         model="gemma2-9b-it"
#     )
    
#     sql_query = response.choices[0].message.content.strip()
#     return sql_query


# def generate_visualization(data, query, chart_type):
 
#     if not data:
#         return None
    

#     keys = data[0].keys()
#     columns = {key: [row[key] for row in data] for key in keys}
    
#     fig = None
    
#     if chart_type == "bar":
     
#         x_col = list(keys)[0]
#         y_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
#         fig = go.Figure(data=[go.Bar(x=columns[x_col], y=columns[y_col])])
#         fig.update_layout(title=query, xaxis_title=x_col, yaxis_title=y_col)
    
#     elif chart_type == "pie":
       
#         label_col = list(keys)[0]
#         value_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
#         fig = go.Figure(data=[go.Pie(labels=columns[label_col], values=columns[value_col])])
#         fig.update_layout(title=query)
    
#     elif chart_type == "line":
       
#         x_col = list(keys)[0]
#         y_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
#         fig = go.Figure(data=[go.Scatter(x=columns[x_col], y=columns[y_col], mode='lines+markers')])
#         fig.update_layout(title=query, xaxis_title=x_col, yaxis_title=y_col)
    
#     elif chart_type == "scatter":
  
#         x_col = list(keys)[0]
#         y_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
#         fig = go.Figure(data=[go.Scatter(x=columns[x_col], y=columns[y_col], mode='markers')])
#         fig.update_layout(title=query, xaxis_title=x_col, yaxis_title=y_col)
    
#     elif chart_type == "histogram":
       
#         data_col = list(keys)[0]
#         fig = go.Figure(data=[go.Histogram(x=columns[data_col])])
#         fig.update_layout(title=query, xaxis_title=data_col, yaxis_title="Count")
    
#     elif chart_type == "heatmap":
#         if len(keys) >= 3:
#             x_col = list(keys)[0]
#             y_col = list(keys)[1]
#             z_col = list(keys)[2]
            
          
#             x_vals = sorted(set(columns[x_col]))
#             y_vals = sorted(set(columns[y_col]))
#             z_matrix = [[0 for _ in range(len(x_vals))] for _ in range(len(y_vals))]
            
#             for i, row in enumerate(data):
#                 x_idx = x_vals.index(row[x_col])
#                 y_idx = y_vals.index(row[y_col])
#                 z_matrix[y_idx][x_idx] = row[z_col]
            
#             fig = go.Figure(data=[go.Heatmap(z=z_matrix, x=x_vals, y=y_vals)])
#             fig.update_layout(title=query)
#         else:
           
#             chart_type = "table"
    
#     if chart_type == "table" or fig is None:
       
#         header_values = list(keys)
#         cell_values = [columns[key] for key in keys]
        
#         fig = go.Figure(data=[go.Table(
#             header=dict(values=header_values),
#             cells=dict(values=cell_values)
#         )])
#         fig.update_layout(title=query)
    
#     return json.dumps(fig, cls=plotly.utils.PlotlyJSONEncoder)



# class Candidate(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     name = db.Column(db.String(100))
#     date = db.Column(db.DateTime, default=datetime.now)
#     email = db.Column(db.String(100))
#     phone_number = db.Column(db.String(20))
#     cv_file = db.Column(db.String(255))
#     current_status = db.Column(db.String(50), default='CV Review')
#     status_due_date = db.Column(db.DateTime, nullable=True)
#     assignee = db.Column(db.String(100), nullable=True)
#     position = db.Column(db.String(100), nullable=True)
#     notified = db.Column(db.Boolean, default=False)
#     fail_stage = db.Column(db.String(50), nullable=True)
#     failed_reason = db.Column(db.Text, nullable=True)
#     source_notes = db.Column(db.Text, nullable=True)
#     last_updated = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)

#     def __repr__(self):
#         return f'<Candidate {self.name}>'


# def extract_text_from_pdf(pdf_path):
#     text = ""
#     with open(pdf_path, 'rb') as file:
#         reader = PyPDF2.PdfReader(file)
#         for page in reader.pages:
#             text += page.extract_text()
#     return text


# def extract_text_from_docx(docx_path):
#     doc = docx.Document(docx_path)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     return text


# def extract_cv_info(text):

#     email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     phone_pattern = r'(\+\d{1,3}[\s-]?)?(\(?\d{3}\)?[\s-]?)?\d{3}[\s-]?\d{4}'
#     name_pattern = r'^([A-Z][a-z]+([\s-][A-Z][a-z]+)+)'
    
#     #n
#     email = re.search(email_pattern, text)
#     phone = re.search(phone_pattern, text)
#     name = re.search(name_pattern, text, re.MULTILINE)
    

#     first_part = text[:100]
    
#     return {
#         'name': name.group(0) if name else None,
#         'email': email.group(0) if email else None,
#         'phone': phone.group(0) if phone else None,
#         'first_part': first_part  
#     }

# # Initialize the database
# with app.app_context():
#     db.create_all()



# # Routes

# @app.route('/')
# def home():
#     return render_template('index.html')

# @app.route('/index')
# def index():
#     candidates = Candidate.query.all()
#     return render_template('cv/index.html', candidates=candidates)

# @app.route('/upload', methods=['GET', 'POST'])
# def upload_cv():
#     if request.method == 'POST':
#         # Check if file part exists
#         if 'file' not in request.files:
#             flash('No file part')
#             return redirect(request.url)
        
#         file = request.files['file']
        
#         # If user does not select file
#         if file.filename == '':
#             flash('No selected file')
#             return redirect(request.url)
        
#         if file:
#             filename = secure_filename(file.filename)
#             file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#             file.save(file_path)
            
#             # Extract text based on file type
#             text = ""
#             if filename.lower().endswith('.pdf'):
#                 text = extract_text_from_pdf(file_path)
#             elif filename.lower().endswith('.docx'):
#                 text = extract_text_from_docx(file_path)
#             else:
#                 flash('Unsupported file format. Please upload PDF or DOCX')
#                 return redirect(request.url)
            
#             # Extract information from CV text
#             info = extract_cv_info(text)
            
#             # Create a new candidate entry
#             new_candidate = Candidate(
#                 name=info['name'],
#                 email=info['email'],
#                 phone_number=info['phone'],
#                 cv_file=filename
#             )
            
#             db.session.add(new_candidate)
#             db.session.commit()
            
#             flash('CV uploaded successfully!')
#             return redirect(url_for('edit_candidate', id=new_candidate.id))
    
#     return render_template('cv/upload.html')

# @app.route('/candidate/<int:id>', methods=['GET', 'POST'])
# def edit_candidate(id):
#     candidate = Candidate.query.get_or_404(id)
    
#     if request.method == 'POST':
#         # Update candidate information from form
#         candidate.name = request.form['name']
#         candidate.email = request.form['email']
#         candidate.phone_number = request.form['phone_number']
#         candidate.current_status = request.form['current_status']
        
#         if request.form['status_due_date']:
#             candidate.status_due_date = datetime.strptime(request.form['status_due_date'], '%Y-%m-%d')
        
#         candidate.assignee = request.form['assignee']
#         candidate.position = request.form['position']
#         candidate.notified = 'notified' in request.form
#         candidate.fail_stage = request.form['fail_stage'] if request.form['fail_stage'] else None
#         candidate.failed_reason = request.form['failed_reason']
#         candidate.source_notes = request.form['source_notes']
        
#         # Update last_updated timestamp
#         candidate.last_updated = datetime.now()
        
#         db.session.commit()
#         flash('Candidate information updated successfully!')
#         return redirect(url_for('index'))
    
#     return render_template('cv/edit.html', candidate=candidate)

# @app.route('/delete/<int:id>', methods=['POST'])
# def delete_candidate(id):
#     candidate = Candidate.query.get_or_404(id)
    
#     # Delete the CV file if it exists
#     if candidate.cv_file:
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], candidate.cv_file)
#         if os.path.exists(file_path):
#             os.remove(file_path)
    
#     db.session.delete(candidate)
#     db.session.commit()
#     flash('Candidate deleted successfully!')
#     return redirect(url_for('index'))

# # Add a simple template filter for date formatting
# @app.template_filter('format_date')
# def format_date(date, format='%Y-%m-%d'):
#     if date:
#         return date.strftime(format)
#     return ''

# @app.route("/index_report")
# def index_report():
#     # Query the database to 
#     conn = get_db_connection()
#     cursor = conn.cursor()
#     cursor.execute("SELECT Current_Status, COUNT(*) as count FROM candidates GROUP BY Current_Status")
#     data = cursor.fetchall()
#     conn.close()
    
#     statuses = [row["Current_Status"] for row in data]
#     counts = [row["count"] for row in data]
    
#     # Create a pie chart using Plotly
#     pie_chart = go.Figure(data=[go.Pie(labels=statuses, values=counts, hole=0.3)])
#     pie_chart.update_layout(title="Candidate Status Distribution")
    
   
#     graphJSON = json.dumps(pie_chart, cls=plotly.utils.PlotlyJSONEncoder)
    
   
#     return render_template('report/index.html', graphJSON=graphJSON)


# @app.route('/generate_report', methods=['POST'])
# def generate_report():

#     conn = get_db_connection()
#     cursor = conn.cursor()
#     cursor.execute("SELECT Current_Status, COUNT(*) as count FROM candidates GROUP BY Current_Status")
#     status_data = cursor.fetchall()
#     conn.close()
    
#     # Prepare a prompt for generating a report
#     report_prompt = "Generate a recruitment report summarizing the candidate status counts:\n"
#     for row in status_data:
#         report_prompt += f"- {row['Current_Status']}: {row['count']}\n"
#     report_prompt += "\nThe report should include insights and recommendations."
    
#     # Generate a report using Groq 
#     chat_completion = groq_client.chat.completions.create(
#         messages=[
#             {
#                 "role": "system", 
#                 "content": system_prompt
#             },
#             {
#                 "role": "user",
#                 "content": report_prompt,
#             }
#         ],
#         model="gemma2-9b-it",
#     )
    
#     report = chat_completion.choices[0].message.content
    
#     # Render the report in a new page
#     return render_template('report/report.html', report=report)

# @app.route('/nl_query')
# def nl_query():
    
#     return render_template('report/nl_query.html')

# @app.route('/query', methods=['POST'])
# def process_query():
   
#     natural_query = request.form.get('query', '')
        
       
#     sql_query = natural_language_to_sql(natural_query)
        
     
#     conn = get_db_connection()
#     cursor = conn.cursor()
#     cursor.execute(sql_query)
#     data = [dict(row) for row in cursor.fetchall()]
#     conn.close()
        
     
#     chart_type = determine_chart_type(natural_query, data)
       
#     graph_json = generate_visualization(data, natural_query, chart_type)
        
#     return jsonify({
#             'status': 'success',
#             'sql_query': sql_query,
#             'data': data,
#             'chart_type': chart_type,
#             'graph_json': graph_json
#         })
 

# if __name__ == '__main__':
#     app.run(debug=True)

import os
import re
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import sqlite3
import plotly.graph_objs as go
import plotly
import json
from groq import Groq
import os
import json
import requests
from typing import Dict, List, Optional
from datetime import datetime, timedelta
from collections import defaultdict
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify
import markdown

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///recruitment.db'  # Changed to match the db name used in get_db_connection
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # Fixed typo in config name
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB max file size

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

groq_client = Groq(
    api_key="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"
)


system_prompt = """
# Recruitment Analysis Assistant

You are an expert AI recruitment analyst specializing in talent acquisition metrics and workforce insights. Your primary role is to analyze recruitment data and generate comprehensive, data-driven reports with strategic recommendations.

## Your Capabilities

- Analyze candidate status distributions and identify bottlenecks in the recruitment pipeline
- Generate actionable insights based on recruitment metrics
- Provide strategic recommendations to improve hiring efficiency
- Identify trends and patterns in recruitment data
- Offer industry benchmarks and best practices when relevant

## Guidelines for Analysis

When analyzing recruitment data:
1. First acknowledge the current state of the recruitment pipeline based on the provided metrics
2. Identify potential issues or bottlenecks in the recruitment process
3. Provide concrete, actionable recommendations with expected outcomes
4. Include quantitative targets where appropriate
5. Consider industry standards and best practices in your analysis
6. Organize insights in a clear, structured format

## Response Format

Structure your reports with the following sections:
- Executive Summary (brief overview of key findings)
- Current Pipeline Status (detailed analysis of provided metrics)
- Key Insights (interpretation of the data and identification of patterns)
- Strategic Recommendations (3-5 specific, actionable items)
- Expected Outcomes (projections for implementing recommendations)

Remember to maintain a professional, consultative tone while providing insights that would be valuable to recruitment managers and HR executives.
"""

# Creating a connection
def get_db_connection():
    conn = sqlite3.connect('recruitment.db')
    conn.row_factory = sqlite3.Row  # to get dict-like rows
    return conn


def get_table_schema():
   
    conn = get_db_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = cursor.fetchall()
    
    schema = {}
    for table in tables:
        table_name = table['name']
        cursor.execute(f"PRAGMA table_info({table_name});")
        columns = cursor.fetchall()
        schema[table_name] = [column['name'] for column in columns]
    
    conn.close()
    return schema


def natural_language_to_sql(query):
    
    schema = get_table_schema()
    
    schema_info = "Database Schema:\n"
    for table, columns in schema.items():
        schema_info += f"Table: {table}\n"
        schema_info += f"Columns: {', '.join(columns)}\n\n"
    
    # Create the prompt
    prompt = f"""
{schema_info}

Based on the schema above, convert the following natural language query to a valid SQL query:
"{query}"

Guidelines:
- Return ONLY the SQL query without any explanation or markdown formatting
- Ensure the query is valid SQLite syntax
- If aggregation is requested, include appropriate GROUP BY statements
- If the query is asking for a visualization, include an appropriate SELECT statement for the visualization
- Return a query that will return a reasonable number of rows (use LIMIT if necessary)
- For time-based queries, use the appropriate SQL date functions
- If you're not sure about a column name, make your best guess based on the schema

SQL Query:
"""
    
    # Generate SQL query
    response = groq_client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": "You are an expert in converting natural language to SQL queries. You carefully analyze database schemas and produce only valid SQL code."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        model="gemma2-9b-it"
    )
    
    sql_query = response.choices[0].message.content.strip()
    return sql_query


# Adding the missing determine_chart_type function
def determine_chart_type(query, data):
    """
    Determine the most appropriate chart type based on the query and result data.
    """
    query = query.lower()
    
    # If no data, return table view
    if not data or len(data) == 0:
        return "table"
    
    # Get number of columns in result
    columns = len(data[0].keys())
    rows = len(data)
    
    # Keywords that suggest chart types
    pie_keywords = ["distribution", "breakdown", "proportion", "percentage", "ratio", "share", "pie"]
    bar_keywords = ["compare", "comparison", "ranking", "rank", "count by", "counts of", "bar chart"]
    line_keywords = ["trend", "over time", "timeline", "progression", "line chart", "history"]
    scatter_keywords = ["correlation", "relationship between", "scatter", "plot"]
    histogram_keywords = ["frequency", "distribution of values", "histogram"]
    heatmap_keywords = ["matrix", "grid", "heatmap", "heat map", "correlation matrix"]
    
    # Check for explicit chart type requests
    for keyword in pie_keywords:
        if keyword in query:
            return "pie"
    
    for keyword in bar_keywords:
        if keyword in query:
            return "bar"
    
    for keyword in line_keywords:
        if keyword in query:
            return "line"
    
    for keyword in scatter_keywords:
        if keyword in query:
            return "scatter"
    
    for keyword in histogram_keywords:
        if keyword in query:
            return "histogram"
    
    for keyword in heatmap_keywords:
        if keyword in query and columns >= 3:
            return "heatmap"
    
    # Default logic based on result structure
    if columns == 2:
        # If one column appears to be categories and the other numeric
        first_col_key = list(data[0].keys())[0]
        second_col_key = list(data[0].keys())[1]
        
        # Check if second column is numeric
        try:
            float(data[0][second_col_key])
            # If few categories, pie chart might be good
            if rows <= 10:
                return "pie"
            else:
                return "bar"
        except (ValueError, TypeError):
            # Both columns might be non-numeric
            return "table"
    
    elif columns >= 3:
        # Multiple data points - table is safest
        return "table"
    
    # Default to table for anything else
    return "table"


def generate_visualization(data, query, chart_type):
 
    if not data:
        return None
    
    keys = data[0].keys()
    columns = {key: [row[key] for row in data] for key in keys}
    
    fig = None
    
    if chart_type == "bar":
     
        x_col = list(keys)[0]
        y_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
        fig = go.Figure(data=[go.Bar(x=columns[x_col], y=columns[y_col])])
        fig.update_layout(title=query, xaxis_title=x_col, yaxis_title=y_col)
    
    elif chart_type == "pie":
       
        label_col = list(keys)[0]
        value_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
        fig = go.Figure(data=[go.Pie(labels=columns[label_col], values=columns[value_col])])
        fig.update_layout(title=query)
    
    elif chart_type == "line":
       
        x_col = list(keys)[0]
        y_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
        fig = go.Figure(data=[go.Scatter(x=columns[x_col], y=columns[y_col], mode='lines+markers')])
        fig.update_layout(title=query, xaxis_title=x_col, yaxis_title=y_col)
    
    elif chart_type == "scatter":
  
        x_col = list(keys)[0]
        y_col = list(keys)[1] if len(keys) > 1 else list(keys)[0]
        fig = go.Figure(data=[go.Scatter(x=columns[x_col], y=columns[y_col], mode='markers')])
        fig.update_layout(title=query, xaxis_title=x_col, yaxis_title=y_col)
    
    elif chart_type == "histogram":
       
        data_col = list(keys)[0]
        fig = go.Figure(data=[go.Histogram(x=columns[data_col])])
        fig.update_layout(title=query, xaxis_title=data_col, yaxis_title="Count")
    
    elif chart_type == "heatmap":
        if len(keys) >= 3:
            x_col = list(keys)[0]
            y_col = list(keys)[1]
            z_col = list(keys)[2]
            
          
            x_vals = sorted(set(columns[x_col]))
            y_vals = sorted(set(columns[y_col]))
            z_matrix = [[0 for _ in range(len(x_vals))] for _ in range(len(y_vals))]
            
            for i, row in enumerate(data):
                x_idx = x_vals.index(row[x_col])
                y_idx = y_vals.index(row[y_col])
                z_matrix[y_idx][x_idx] = row[z_col]
            
            fig = go.Figure(data=[go.Heatmap(z=z_matrix, x=x_vals, y=y_vals)])
            fig.update_layout(title=query)
        else:
           
            chart_type = "table"
    
    if chart_type == "table" or fig is None:
       
        header_values = list(keys)
        cell_values = [columns[key] for key in keys]
        
        fig = go.Figure(data=[go.Table(
            header=dict(values=header_values),
            cells=dict(values=cell_values)
        )])
        fig.update_layout(title=query)
    
    return json.dumps(fig, cls=plotly.utils.PlotlyJSONEncoder)


# Changed Candidate model table name to match database
class Candidate(db.Model):
    __tablename__ = 'candidates'  # Added tablename to match SQL queries
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))
    date = db.Column(db.DateTime, default=datetime.now)
    email = db.Column(db.String(100))
    phone_number = db.Column(db.String(20))
    cv_file = db.Column(db.String(255))
    current_status = db.Column(db.String(50), default='CV Review')
    status_due_date = db.Column(db.DateTime, nullable=True)
    assignee = db.Column(db.String(100), nullable=True)
    position = db.Column(db.String(100), nullable=True)
    notified = db.Column(db.Boolean, default=False)
    fail_stage = db.Column(db.String(50), nullable=True)
    failed_reason = db.Column(db.Text, nullable=True)
    source_notes = db.Column(db.Text, nullable=True)
    last_updated = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)

    def __repr__(self):
        return f'<Candidate {self.name}>'


def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text


def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_cv_info(text):

    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'(\+\d{1,3}[\s-]?)?(\(?\d{3}\)?[\s-]?)?\d{3}[\s-]?\d{4}'
    name_pattern = r'^([A-Z][a-z]+([\s-][A-Z][a-z]+)+)'
    
    email = re.search(email_pattern, text)
    phone = re.search(phone_pattern, text)
    name = re.search(name_pattern, text, re.MULTILINE)
    
    first_part = text[:100]
    
    return {
        'name': name.group(0) if name else None,
        'email': email.group(0) if email else None,
        'phone': phone.group(0) if phone else None,
        'first_part': first_part  
    }

class ClickUpManager:
    def __init__(self, api_token: str):
        self.api_token = api_token
        self.headers = {
            "Authorization": api_token,
            "Content-Type": "application/json"
        }
        self.base_url = "https://api.clickup.com/api/v2"

    def get_all_teams(self) -> List[Dict]:
        """Get all teams (workspaces) the user has access to"""
        response = requests.get(
            f"{self.base_url}/team",
            headers=self.headers
        )
        response.raise_for_status()
        return response.json()["teams"]

    def get_spaces_in_team(self, team_id: str) -> List[Dict]:
        """Get all spaces within a specific team"""
        response = requests.get(
            f"{self.base_url}/team/{team_id}/space",
            headers=self.headers
        )
        response.raise_for_status()
        return response.json()["spaces"]

    def get_lists_in_space(self, space_id: str) -> List[Dict]:
        """Get all lists within a space"""
        url = f"{self.base_url}/space/{space_id}/list"
        response = requests.get(url, headers=self.headers)
        response.raise_for_status()
        return response.json()["lists"]

    def get_folder_lists(self, folder_id: str) -> List[Dict]:
        """Get all lists within a folder"""
        url = f"{self.base_url}/folder/{folder_id}/list"
        response = requests.get(url, headers=self.headers)
        response.raise_for_status()
        return response.json()["lists"]

    def get_folders_in_space(self, space_id: str) -> List[Dict]:
        """Get all folders within a space"""
        url = f"{self.base_url}/space/{space_id}/folder"
        response = requests.get(url, headers=self.headers)
        response.raise_for_status()
        return response.json()["folders"]

    def get_tasks_in_list(self, list_id: str, params: Optional[Dict] = None) -> List[Dict]:
        """Get all tasks within a list"""
        url = f"{self.base_url}/list/{list_id}/task"
        response = requests.get(url, headers=self.headers, params=params)
        response.raise_for_status()
        return response.json()["tasks"]

    def get_space_details(self, space_id: str) -> Dict:
        """Get space information"""
        response = requests.get(
            f"{self.base_url}/space/{space_id}",
            headers=self.headers
        )
        response.raise_for_status()
        return response.json()

class ClickUpTaskCounter(ClickUpManager):
    def count_tasks_in_space(self, space_id: str, days_back: Optional[int] = None) -> Dict:
        """
        Count tasks in a space with detailed breakdown

        Args:
            space_id (str): ID of the space to analyze
            days_back (int, optional): If provided, only count tasks from the last X days
        """
        # Initialize counters
        task_stats = {
            "total_tasks": 0,
            "completed_tasks": 0,
            "open_tasks": 0,
            "tasks_by_status": {},
            "tasks_by_priority": {
                "urgent": 0,
                "high": 0,
                "normal": 0,
                "low": 0,
                "no_priority": 0
            },
            "lists_count": 0,
            "folders_count": 0
        }

        # Set up date filtering if specified
        params = {}
        if days_back:
            start_date = datetime.now() - timedelta(days=days_back)
            params["date_created_gt"] = int(start_date.timestamp() * 1000)

        try:
            # Get folders in space
            folders = self.get_folders_in_space(space_id)
            task_stats["folders_count"] = len(folders)

            # Process folderless lists
            space_lists = self.get_lists_in_space(space_id)
            all_lists = space_lists.copy()

            # Process folders and their lists
            for folder in folders:
                folder_lists = self.get_folder_lists(folder["id"])
                all_lists.extend(folder_lists)

            task_stats["lists_count"] = len(all_lists)

            # Process all lists
            for list_item in all_lists:
                tasks = self.get_tasks_in_list(list_item["id"], params)

                for task in tasks:
                    task_stats["total_tasks"] += 1

                    # Count by status
                    status = task["status"]["status"]
                    task_stats["tasks_by_status"][status] = task_stats["tasks_by_status"].get(status, 0) + 1

                    if status.lower() in ["complete", "completed", "done"]:
                        task_stats["completed_tasks"] += 1
                    else:
                        task_stats["open_tasks"] += 1

                    # Count by priority
                    priority = task.get("priority")
                    if priority:
                        priority_name = priority["priority"].lower()
                        task_stats["tasks_by_priority"][priority_name] += 1
                    else:
                        task_stats["tasks_by_priority"]["no_priority"] += 1

            return task_stats
        except Exception as e:
            print(f"Error counting tasks: {e}")
            return task_stats

class SpaceAssigneeTracker(ClickUpManager):
    def get_space_assignees(self, space_id: str) -> Dict:
        """
        Get all assignees and their tasks in a specific space
        """
        # Initialize data structure for assignees
        assignee_data = defaultdict(lambda: {
            "name": "",
            "email": "",
            "username": "",
            "task_count": 0,
            "tasks": [],
            "lists": set()
        })

        try:
            # Get space details
            space = self.get_space_details(space_id)

            # Get all lists in the space
            lists = self.get_lists_in_space(space_id)

            for list_item in lists:
                tasks = self.get_tasks_in_list(list_item['id'])

                for task in tasks:
                    for assignee in task.get("assignees", []):
                        assignee_id = assignee["id"]

                        # Update assignee information
                        assignee_data[assignee_id].update({
                            "name": assignee.get("username", "No username"),
                            "email": assignee.get("email", "No email"),
                            "username": assignee.get("username", "No username")
                        })

                        # Update task information
                        assignee_data[assignee_id]["task_count"] += 1
                        assignee_data[assignee_id]["lists"].add(list_item["name"])

                        # Add task details
                        task_info = {
                            "task_id": task["id"],
                            "task_name": task["name"],
                            "status": task["status"]["status"],
                            "due_date": task.get("due_date", "No due date"),
                            "list_name": list_item["name"],
                            "priority": task.get("priority", {}).get("priority", "No priority")
                        }
                        assignee_data[assignee_id]["tasks"].append(task_info)

            # Convert sets to lists for JSON serialization
            for assignee_id in assignee_data:
                assignee_data[assignee_id]["lists"] = list(assignee_data[assignee_id]["lists"])

            return dict(assignee_data)
        except Exception as e:
            print(f"Error retrieving assignees: {e}")
            return {}

class ReportGenerator:
    def __init__(self, api_key=None):
        """Initialize with optional API key for GPT integration"""
        self.api_key = api_key
        if api_key:
            self.api_url = "https://api.groq.com/openai/v1/chat/completions"
            self.headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
        
    def _prepare_data_summary(self, data: Dict) -> Dict:
        """Prepare a summary of the data for the LLM"""
        summary = {
            "total_assignees": len(data),
            "total_tasks": sum(assignee["task_count"] for assignee in data.values()),
            "assignee_summaries": []
        }

        for assignee_id, assignee_data in data.items():
            # Calculate task status distribution
            status_count = {}
            for task in assignee_data["tasks"]:
                status = task["status"]
                status_count[status] = status_count.get(status, 0) + 1

            # Calculate priority distribution
            priority_count = {}
            for task in assignee_data["tasks"]:
                priority = task.get("priority", "No priority")
                priority_count[priority] = priority_count.get(priority, 0) + 1

            summary["assignee_summaries"].append({
                "name": assignee_data["name"],
                "email": assignee_data["email"],
                "task_count": assignee_data["task_count"],
                "lists": assignee_data["lists"],
                "status_distribution": status_count,
                "priority_distribution": priority_count
            })

        return summary

    def generate_report(self, assignee_data: Dict, task_stats: Dict, space_name: str) -> str:
        """Generate a report combining assignee data and task statistics"""
        if not self.api_key:
            # If no API key, generate a basic report
            return self._generate_basic_report(assignee_data, task_stats, space_name)
        
        # Prepare data summary for LLM
        summary = self._prepare_data_summary(assignee_data)
        
        # Create prompt for LLM
        prompt = f"""
        Please analyze this ClickUp workspace data and create a comprehensive report. 

        Workspace: {space_name}
        
        Task Statistics:
        Total Tasks: {task_stats['total_tasks']}
        Completed Tasks: {task_stats['completed_tasks']}
        Open Tasks: {task_stats['open_tasks']}
        Number of Lists: {task_stats['lists_count']}
        Number of Folders: {task_stats['folders_count']}
        
        Tasks by Status: {json.dumps(task_stats['tasks_by_status'])}
        Tasks by Priority: {json.dumps(task_stats['tasks_by_priority'])}
        
        Assignee Information:
        Total Assignees: {summary['total_assignees']}
        
        Detailed Assignee Information:
        {json.dumps(summary['assignee_summaries'], indent=2)}

        Please create a professional report that includes:
        1. Executive Summary
        2. Workload Distribution Analysis
        3. Task Status Overview
        4. Priority Distribution Analysis
        5. Team Member Performance Insights
        6. Recommendations for Workload Balancing
        7. Potential Bottlenecks or Areas of Concern

        Make the report data-driven but easy to understand. Include specific numbers and percentages where relevant.
        Format the report in Markdown.
        """

        try:
            # Generate report using GROQ
            payload = {
                "model": "mixtral-8x7b-32768",  # Using Mixtral model
                "messages": [
                    {"role": "system", "content": "You are a professional project management analyst creating a report based on ClickUp workspace data."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 2000
            }

            response = requests.post(
                self.api_url,
                headers=self.headers,
                json=payload
            )

            response.raise_for_status()
            report = response.json()["choices"][0]["message"]["content"]
            return report

        except Exception as e:
            print(f"Error generating report with LLM: {str(e)}")
            # Fall back to basic report
            return self._generate_basic_report(assignee_data, task_stats, space_name)

    def _generate_basic_report(self, assignee_data: Dict, task_stats: Dict, space_name: str) -> str:
        """Generate a basic report without using LLM"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        report = f"""# ClickUp Workspace Analysis Report
        
## Workspace: {space_name}
Generated on: {timestamp}

## Executive Summary
This report provides an analysis of your ClickUp workspace "{space_name}".

## Task Statistics
- **Total Tasks**: {task_stats['total_tasks']}
- **Completed Tasks**: {task_stats['completed_tasks']} ({task_stats['completed_tasks']/task_stats['total_tasks']*100:.1f}% if tasks exist else 0)
- **Open Tasks**: {task_stats['open_tasks']}
- **Number of Lists**: {task_stats['lists_count']}
- **Number of Folders**: {task_stats['folders_count']}

## Task Status Distribution
"""
        
        for status, count in task_stats['tasks_by_status'].items():
            percentage = (count / task_stats['total_tasks'] * 100) if task_stats['total_tasks'] > 0 else 0
            report += f"- **{status}**: {count} ({percentage:.1f}%)\n"
            
        report += """
## Task Priority Distribution
"""
        
        for priority, count in task_stats['tasks_by_priority'].items():
            percentage = (count / task_stats['total_tasks'] * 100) if task_stats['total_tasks'] > 0 else 0
            report += f"- **{priority.capitalize()}**: {count} ({percentage:.1f}%)\n"
            
        report += """
## Assignee Workload
"""
        
        for assignee_id, data in assignee_data.items():
            report += f"""
### {data['name']}
- **Email**: {data['email']}
- **Total Tasks**: {data['task_count']}
- **Active in Lists**: {', '.join(data['lists'][:5])}{"..." if len(data['lists']) > 5 else ""}
"""
            
        report += """
## Recommendations
1. Review workload distribution among team members to ensure balanced assignments
2. Address any tasks with high priority that remain unresolved
3. Consider consolidating or archiving unused lists to streamline workspace


"""
        
        return report


# Initialize the database
with app.app_context():
    db.create_all()


# Routes

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/index')
def index():
    candidates = Candidate.query.all()
    return render_template('cv/index.html', candidates=candidates)

@app.route('/upload', methods=['GET', 'POST'])
def upload_cv():
    if request.method == 'POST':
        # Check if file part exists
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        
        file = request.files['file']
        
        # If user does not select file
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        
        if file:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Extract text based on file type
            text = ""
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                flash('Unsupported file format. Please upload PDF or DOCX')
                return redirect(request.url)
            
            # Extract information from CV text
            info = extract_cv_info(text)
            
            # Create a new candidate entry
            new_candidate = Candidate(
                name=info['name'],
                email=info['email'],
                phone_number=info['phone'],
                cv_file=filename
            )
            
            db.session.add(new_candidate)
            db.session.commit()
            
            flash('CV uploaded successfully!')
            return redirect(url_for('edit_candidate', id=new_candidate.id))
    
    return render_template('cv/upload.html')

@app.route('/candidate/<int:id>', methods=['GET', 'POST'])
def edit_candidate(id):
    candidate = Candidate.query.get_or_404(id)
    
    if request.method == 'POST':
        # Update candidate information from form
        candidate.name = request.form['name']
        candidate.email = request.form['email']
        candidate.phone_number = request.form['phone_number']
        candidate.current_status = request.form['current_status']
        
        if request.form['status_due_date']:
            candidate.status_due_date = datetime.strptime(request.form['status_due_date'], '%Y-%m-%d')
        
        candidate.assignee = request.form['assignee']
        candidate.position = request.form['position']
        candidate.notified = 'notified' in request.form
        candidate.fail_stage = request.form['fail_stage'] if request.form['fail_stage'] else None
        candidate.failed_reason = request.form['failed_reason']
        candidate.source_notes = request.form['source_notes']
        
        # Update last_updated timestamp
        candidate.last_updated = datetime.now()
        
        db.session.commit()
        flash('Candidate information updated successfully!')
        return redirect(url_for('index'))
    
    return render_template('cv/edit.html', candidate=candidate)

@app.route('/delete/<int:id>', methods=['POST'])
def delete_candidate(id):
    candidate = Candidate.query.get_or_404(id)
    
    # Delete the CV file if it exists
    if candidate.cv_file:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], candidate.cv_file)
        if os.path.exists(file_path):
            os.remove(file_path)
    
    db.session.delete(candidate)
    db.session.commit()
    flash('Candidate deleted successfully!')
    return redirect(url_for('index'))

# Add a simple template filter for date formatting
@app.template_filter('format_date')
def format_date(date, format='%Y-%m-%d'):
    if date:
        return date.strftime(format)
    return ''

@app.route("/index_report")
def index_report():
    # Query the database
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT current_status AS Current_Status, COUNT(*) as count FROM candidates GROUP BY current_status")
    data = cursor.fetchall()
    conn.close()
    
    statuses = [row["Current_Status"] for row in data]
    counts = [row["count"] for row in data]
    
    # Create a pie chart using Plotly
    pie_chart = go.Figure(data=[go.Pie(labels=statuses, values=counts, hole=0.3)])
    pie_chart.update_layout(title="Candidate Status Distribution")
    
    graphJSON = json.dumps(pie_chart, cls=plotly.utils.PlotlyJSONEncoder)
    
    return render_template('report/index.html', graphJSON=graphJSON)


@app.route('/generate_report', methods=['POST'])
def generate_report():

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT current_status AS Current_Status, COUNT(*) as count FROM candidates GROUP BY current_status")
    status_data = cursor.fetchall()
    conn.close()
    
    # Prepare a prompt for generating a report
    report_prompt = "Generate a recruitment report summarizing the candidate status counts:\n"
    for row in status_data:
        report_prompt += f"- {row['Current_Status']}: {row['count']}\n"
    report_prompt += "\nThe report should include insights and recommendations."
    
    # Generate a report using Groq 
    chat_completion = groq_client.chat.completions.create(
        messages=[
            {
                "role": "system", 
                "content": system_prompt
            },
            {
                "role": "user",
                "content": report_prompt,
            }
        ],
        model="gemma2-9b-it",
    )
    
    report = chat_completion.choices[0].message.content
    
    # Render the report in a new page
    return render_template('report/report.html', report=report)

@app.route('/nl_query')
def nl_query():
    return render_template('report/nl_query.html')

@app.route('/query', methods=['POST'])
def process_query():
    natural_query = request.form.get('query', '')
    
    sql_query = natural_language_to_sql(natural_query)
    
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(sql_query)
    data = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    chart_type = determine_chart_type(natural_query, data)
    
    graph_json = generate_visualization(data, natural_query, chart_type)
    
    return jsonify({
        'status': 'success',
        'sql_query': sql_query,
        'data': data,
        'chart_type': chart_type,
        'graph_json': graph_json
    })
 

@app.route('/index_clickup')
def index_clickup():
    api_token = session.get('api_token')
    if not api_token:
        return render_template('clickup/login.html')
    return render_template(url_for('workspaces'))

@app.route('/login', methods=['POST'])
def login():
    api_token = request.form.get('api_token')
    
    if not api_token:
        flash('API token is required', 'danger')
        return redirect(url_for('index_clickup'))
    
    # Verify API token works
    try:
        manager = ClickUpManager(api_token)
        manager.get_all_teams()
        session['api_token'] = api_token
        flash('Successfully logged in', 'success')
        return redirect(url_for('workspaces'))
    except Exception as e:
        flash(f'Invalid API token: {str(e)}', 'danger')
        return redirect(url_for('index_clickup'))

@app.route('/logout')
def logout():
    session.pop('api_token', None)
    flash('Logged out successfully', 'success')
    return redirect(url_for('index_clickup'))


@app.route('/workspaces')
def workspaces():
    api_token = session.get('api_token')
    if not api_token:
        flash('Please log in first', 'warning')
        return redirect(url_for('index_clickup'))
    
    try:
        manager = ClickUpManager(api_token)
        teams = manager.get_all_teams()
        return render_template('clickup/workspaces.html', teams=teams)
    except Exception as e:
        flash(f'Error retrieving workspaces: {str(e)}', 'danger')
        return redirect(url_for('index_clickup'))

@app.route('/spaces/<team_id>')
def spaces(team_id):
    api_token = session.get('api_token')
    if not api_token:
        flash('Please log in first', 'warning')
        return redirect(url_for('index_clickup'))
    
    try:
        manager = ClickUpManager(api_token)
        spaces = manager.get_spaces_in_team(team_id)
        return render_template('clickup/spaces.html', spaces=spaces, team_id=team_id)
    except Exception as e:
        flash(f'Error retrieving spaces: {str(e)}', 'danger')
        return redirect(url_for('workspaces'))

@app.route('/space/<space_id>')
def space_dashboard(space_id):
    api_token = session.get('api_token')
    if not api_token:
        flash('Please log in first', 'warning')
        return redirect(url_for('index_clickup'))
    
    days_back = request.args.get('days_back', default=30, type=int)
    
    try:
        manager = ClickUpManager(api_token)
        space_details = manager.get_space_details(space_id)
        
        # Get task statistics
        task_counter = ClickUpTaskCounter(api_token)
        task_stats = task_counter.count_tasks_in_space(space_id, days_back=days_back)
        
        # Get assignee data
        assignee_tracker = SpaceAssigneeTracker(api_token)
        assignee_data = assignee_tracker.get_space_assignees(space_id)
        
        return render_template(
            'clickup/space_dashboard.html', 
            space=space_details, 
            task_stats=task_stats,
            assignee_data=assignee_data,
            days_back=days_back
        )
    except Exception as e:
        flash(f'Error retrieving space data: {str(e)}', 'danger')
        return redirect(url_for('workspaces'))

@app.route('/generate_report/<space_id>', methods=['POST'])
def generate_report(space_id):
    api_token = session.get('api_token')
    if not api_token:
        flash('Please log in first', 'warning')
        return redirect(url_for('index_clickup'))
    
    groq_api_key = request.form.get('groq_api_key', '')
    
    try:
        # Get space details first
        manager = ClickUpManager(api_token)
        space_details = manager.get_space_details(space_id)
        space_name = space_details.get('name', 'Unknown Space')
        
        # Get task statistics
        task_counter = ClickUpTaskCounter(api_token)
        task_stats = task_counter.count_tasks_in_space(space_id)
        
        # Get assignee data
        assignee_tracker = SpaceAssigneeTracker(api_token)
        assignee_data = assignee_tracker.get_space_assignees(space_id)
        
        # Generate report
        report_generator = ReportGenerator(groq_api_key if groq_api_key else None)
        report_content = report_generator.generate_report(assignee_data, task_stats, space_name)
        
        # Save report to file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"clickup_report_{space_id}_{timestamp}.md"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        with open(filepath, 'w') as f:
            f.write(report_content)
        
        # Store report in session for display
        session['report_content'] = report_content
        session['report_filename'] = filename
        
        flash('Report generated successfully', 'success')
        return redirect(url_for('view_report'))
    except Exception as e:
        flash(f'Error generating report: {str(e)}', 'danger')
        return redirect(url_for('space_dashboard', space_id=space_id))


# Add to your imports
import markdown

@app.route('/view_report')
def view_report():
    api_token = session.get('api_token')
    if not api_token:
        flash('Please log in first', 'warning')
        return redirect(url_for('index_clickup'))
    
    report_content = session.get('report_content')
    if not report_content:
        flash('No report found', 'warning')
        return redirect(url_for('workspaces'))
    
    # Convert markdown to HTML
    html_content = markdown.markdown(report_content)
    
    filename = session.get('report_filename', 'report.md')
    return render_template('clickup/view_report.html', report_content=html_content, filename=filename)

@app.route('/download_report/<filename>')
def download_report(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# API routes for AJAX calls
@app.route('/api/task_stats/<space_id>')
def api_task_stats(space_id):
    api_token = session.get('api_token')
    if not api_token:
        return jsonify({'error': 'Unauthorized'}), 401
    
    days_back = request.args.get('days_back', default=30, type=int)
    
    try:
        # Get task statistics
        task_counter = ClickUpTaskCounter(api_token)
        task_stats = task_counter.count_tasks_in_space(space_id, days_back=days_back)
        return jsonify(task_stats)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)